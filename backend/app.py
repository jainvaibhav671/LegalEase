import logging
import os
import re
from datetime import datetime

import docx
import google.generativeai as genai
import markdown
# from pypdf import PdfReader
import pytesseract
from dotenv import find_dotenv, load_dotenv
from flask import Flask, jsonify, request, send_file, session
from flask_cors import CORS
from flask_session import Session
from PIL import Image
from PyPDF2 import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import inch, landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.platypus.tables import Table, TableStyle
from werkzeug.utils import secure_filename

# from nlp_summary import generate_summary_gpu
# from translate import english_to_hindi


genai.configure(api_key=os.getenv("API_KEY"))
model = genai.GenerativeModel("gemini-pro")
chat = model.start_chat()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("HELLO WORLD")

input_filedir = "pdfs/"
UPLOAD_FOLDER = "./New"
ALLOWED_EXTENSIONS = set(["txt", "pdf", "png", "jpg", "jpeg", "gif"])

app = Flask(__name__)
app.secret_key = "arsdkfja"
SESSION_TYPE = "filesystem"
app.config.from_object(__name__)
Session(app)
CORS(app)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
CORS(app, expose_headers="Authorization")


class readFile:
    # Allowed FileTypes: pdf, docx, jpg, png, txt

    def __init__(self, filepath):
        self.filepath = filepath

    def read_file(self):
        file_extension = self.filepath.split(".")[-1].lower()

        if file_extension == "pdf":
            return self.read_pdf()
        elif file_extension == "jpg":
            return self.read_jpg()
        elif file_extension == "docx":
            return self.read_docx()
        elif file_extension == "txt":
            return self.read_txt()
        else:
            print("Unsupported file type.")

    def read_pdf(self):
        pdf = PdfReader(self.filepath)
        global extracted_text
        extracted_text = ""
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            text = page.extract_text()
            extracted_text += "\n" + text

        return extracted_text

    def read_docx(self):
        doc = docx.Document(self.filepath)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        text = "\n".join(fullText)
        return text

    def read_jpg(self):
        img = Image.open(self.filepath)
        text = pytesseract.image_to_string(img)
        return text

    def read_txt(self):
        with open(self.filepath, "r", encoding="utf-8") as file:
            text = file.read()
        return text


_ = load_dotenv(find_dotenv())


def get_completion(prompt):
    messages = [{"role": "user", "content": prompt}]
    # response = openai.ChatCompletion.create(
    #     model=model,
    #     messages=messages,
    #     temperature=0,
    # )
    # return response.choices[0].message["content"]

    res = model.generate_content(prompt)
    return res.text


class summary:

    def __init__(self, document):
        self.document = document

    def test(self, hindi=False):

        prompt = f"""
        Convert the given text which will contain different types of legal documents such as contracts , purchasing , hiring , deed , general agreement , trademark agreement , leese , rent etc .,to a simple easily human readable format text which is going to be easily understood by a person who has no idea about any law terms and it's language.
        Replace complex legal terminology with plain everyday language , also reduce the complex legal document into short and readable paragraphs which should be easy to understand.
        Text should be in Plain Language Summary format. Also, Simplify the Legal Concept present in the text and avoid ambiguity.
        Generate a title from the given text that whaould be well suited according to the legal document and place that title at top.
        Always include certain details like "document_id", "document_name", "document_type", "creation date", "names of the parties involved", "author", "client_name", "case_number", make seperate enteries for each of these.
        Write the text in form of paragraphs with headings like "Case description"(Make sub headings like "document_id", "document_name", "document_type", "creation date", "names of the parties involved", "author", "client_name", "case_number", if present, and write the entries in front of these), "Clauses"(rights, responsibilities, and terms of the parties involved), "Issues addressed", "references"
        Also, add appropriate line breaks.
        Minimum length of the dtext should be 150
        """

        # response = openai.ChatCompletion.create(
        #     model="gpt-3.5-turbo",
        #     messages=[
        #         {"role": "user", "content": prompt}],
        #     temperature=0.0,
        # )
        # result = response.choices[0]["message"]["content"]
        # return result.replace('"', "'")

        res = model.generate_content(prompt)

        if hindi:
            return english_to_hindi(markdown.markdown(res.text))

        return markdown.markdown(res.text)


def text_to_pdf(text):
    # Define custom styles
    heading_style = ParagraphStyle(
        name="HeadingStyle",
        fontSize=16,
        textColor=colors.blue,
        alignment=1,  # Center alignment
    )

    subheading_style = ParagraphStyle(
        name="SubheadingStyle",
        fontSize=12,
        textColor=colors.red,
        alignment=0,  # Left alignment
    )

    title_style = ParagraphStyle(
        name="TitleStyle",
        fontSize=16,
        textColor=colors.black,
        alignment=1,  # Center alignment
        fontName="Helvetica-Bold",
        leading=18,
    )

    bold_style = getSampleStyleSheet()["Normal"]
    bold_style.fontName = "Helvetica"

    # Create a SimpleDocTemplate to handle word wrapping
    global pdf_filename
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    pdf_filename = f"output_{timestamp}.pdf"

    # Create a SimpleDocTemplate to handle word wrapping
    doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
    line_break_words = [
        "Clauses",
        "Case Description",
        "Issues Addressed",
        "References",
        "Note",
    ]

    # Create a list to hold the flowable elements (e.g., paragraphs)
    story = []

    # Split the input text into paragraphs based on empty lines
    paragraphs = re.split(r"\n\s*\n", text.strip())

    for paragraph in paragraphs:
        # Find and style specific words within the paragraph
        paragraph = re.sub(
            r"\b(Case Description|Clauses|Issues Addressed|References|Note|Signed by)\b",
            r'<b><font size="13">\1</font></b>',
            paragraph,
        )
        paragraph = re.sub(
            r"\b(Document ID|Document Name|Document Type|Creation Date|Names of the Parties Involved|Author|Client Name|Case Number)\b",
            r"<b><i>\1</i></b>",
            paragraph,
        )

        if paragraph.startswith("Title:"):
            # Apply title_style to the title and move to the next line
            p = Paragraph(paragraph.strip().replace("Title:", ""), title_style)
            story.append(p)
            story.append(Spacer(1, 12))
        else:
            # Apply bold_style to other paragraphs
            p = Paragraph(paragraph, bold_style)
            story.append(p)
            story.append(Spacer(1, 12))
        # Create a Paragraph with word wrap
        # p = Paragraph(paragraph, getSampleStyleSheet()['Normal'])
        # story.append(p)
        # story.append(Spacer(1, 12))  # Add space between paragraphs

    doc.build(story)
    return pdf_filename


@app.route("/upload", methods=["POST"])
def fileUpload():

    target = os.path.join(UPLOAD_FOLDER)
    if not os.path.isdir(target):
        os.mkdir(target)
    logger.info("welcome to upload`")
    file = request.files["file"]
    filename = secure_filename(file.filename)
    destination = "/".join([target, filename])
    file.save(destination)
    response = "Whatever you wish too return"
    text = readFile(destination).read_file()

    len(text)

    result = summary(text).test()

    len(result)

    text_to_pdf(result)
    return response


@app.route("/get-pdf", methods=["GET"])
def get_pdf():
    # Replace 'path_to_your_pdf.pdf' with the actual path to your PDF file
    # return send_file(f"./{pdf_filename}", as_attachment=True)

    filename = request.json.get("filename", None)
    if not filename:
        return "No filename was provided", 400

    file_path = os.path.join(input_filedir, filename)
    return send_file(file_path, as_attachment=True)


@app.route("/get_response", methods=["GET", "POST"])
def get_response():

    user_input = request.json["input"]
    return [get_completion(user_input)]


class generate:

    def __init__(self, row_values):
        self.row_values = row_values

    def document(self):

        prompt = f"""
        You are given a input list which consists values in the dictionary form of
        {self.row_values}
        . These values are given to create a legal document out of it.
        Through the given information, form text for a legal document.
        The language used in the document should be the language used by the law firms. I want you to give me a detailed legal document based like this.
        The text should be in the following order:
        Title: The document usually begins with a clear and descriptive title that identifies the type of document and the parties involved.
        Introduction/Preamble: Brief introduction. May include the date of execution and a statement of the parties' intentions.
        Recitals/Consideration: Some legal documents, especially contracts, include a section that sets out the background or reasons for the agreement. This section may also describe the consideration or benefit that each party is receiving from the agreement.
        Definitions: Only In longer and more complex documents, there may be a section that defines key terms used throughout the document.
        Operative Clauses/Agreement: This is the core of the document, where the parties set out their rights, obligations, and responsibilities. In a contract, for example, this section would outline the terms and conditions of the agreement.
        Miscellaneous Clauses: Depending on the document, there may be various additional clauses, such as choice of law, dispute resolution, confidentiality, or indemnification clauses. These clauses address specific issues relevant to the document.
        Termination Conditions:
        Remedies:
        """

        # response = openai.ChatCompletion.create(
        #     model="gpt-3.5-turbo",
        #     messages=[
        #         {"role": "user", "content": prompt}],
        #     temperature=0.0,
        # )
        # result = response.choices[0]["message"]["content"]
        # return result.replace('"', "'")

        res = model.generate_content(prompt)
        return markdown.markdown(res.text)


def style_text(text):
    # Define custom styles
    title_style = ParagraphStyle(
        name="TitleStyle",
        fontSize=18,
        textColor=colors.black,
        alignment=1,  # Center alignment
        fontName="Helvetica-Bold",
    )

    bold_style = getSampleStyleSheet()["Normal"]
    bold_style.fontName = "Helvetica"

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    pdf_filename = f"output_{timestamp}.pdf"

    # Create a SimpleDocTemplate to handle word wrapping
    doc = SimpleDocTemplate(pdf_filename, pagesize=letter)

    # Create a list to hold the flowable elements (e.g., paragraphs)
    story = []

    # Split the input text into paragraphs based on empty lines
    paragraphs = re.split(r"\n\s*\n", text.strip())

    for paragraph in paragraphs:
        # Find and style specific words within the paragraph
        paragraph = re.sub(
            r"\b(Introduction/Preamble|Recitals/Consideration|Definitions|Miscellaneous Clauses|Termination Conditions|Remedies|Operative Clauses/Agreement)\b",
            r"<b>\1</b>",
            paragraph,
        )

        # Create a Paragraph with word wrap and apply the appropriate style
        if paragraph.startswith("Title:"):
            # Apply title_style to the title and move to the next line
            p = Paragraph(paragraph.strip().replace("Title:", ""), title_style)
            story.append(p)
            story.append(Spacer(1, 16))  # Add space between paragraphs
        else:
            # Apply bold_style to other paragraphs
            p = Paragraph(paragraph, bold_style)
            story.append(p)
            story.append(Spacer(1, 12))  # Add space between paragraphs

    doc.build(story)


@app.route("/api/receive-json", methods=["POST"])
def receive_json():
    data = request.json  # JSON data sent from React
    print(data)  # You can process or log the data as needed
    text = generate(data).document()
    style_text(text)
    # result=generate(data).document
    # style_text(result)
    return jsonify({"message": "JSON data received successfully"})


# @app.route("/summarize", methods=["POST"])
# def nlp_summarize():
# input_filename = "file"

# # Create the directory for temporarily save the pdfs
# os.makedirs(input_filedir, exist_ok=True)

# print(request.files)

# if input_filename not in request.files:
#     return "No file found", 400

# # uploaded file
# file = request.files[input_filename]

# if len(file.filename) == 0:
#     return "No file selected for uploading", 400

# file_path = os.path.join(input_filedir, file.filename)
# file.save(file_path)

# file_reader = readFile(file_path)
# extracted_text = file_reader.read_file()

# with open("summary.txt", "w") as f:
#     f.write(extracted_text)

# summary = generate_summary_gpu(extracted_text)

# pdf_filename = text_to_pdf(summary)

# return { "filename": pdf_filename }


@app.route("/summarize", methods=["POST"])
def summarize():
    input_filename = "file"

    # Create the directory for temporarily save the pdfs
    os.makedirs(input_filedir, exist_ok=True)
    print(request.files)

    print("hello")
    if input_filename not in request.files:
        return "No file found", 400

    # uploaded file
    file = request.files[input_filename]

    if len(file.filename) == 0:
        return "No file selected for uploading", 400

    print("File received")
    file_path = os.path.join(input_filedir, file.filename)
    file.save(file_path)

    file_reader = readFile(file_path)
    extracted_text = file_reader.read_file()

    global summary_text

    summary_text = summary(extracted_text).test()
    # summary_text = summary(extracted_text).test(hindi=request.json.get("hindi", False)

    summary_text = summary_text.replace("</p>", "</p><br />")
    summary_text = summary_text.replace("</strong>", "</strong><br />")
    summary_text = summary_text.replace("<p>", "<br /><p>")

    session["original"] = extracted_text
    session["summary"] = summary_text

    pdf_filename = text_to_pdf(summary_text)
    return send_file(pdf_filename)


@app.route("/generate-document", methods=["POST"])
def generate_document():
    jsonData = request.json.get("data")

    text = generate(jsonData).document()

    res = model.generate_content(text)
    filename = text_to_pdf(res.text)

    return send_file(
        open(filename, "rb"), as_attachment=True, download_name=filename, max_age=50000
    )


@app.route("/chat", methods=["POST"])
def chat():

    user_prompt = request.json.get("prompt")

    res = model.generate_content(
        f"""
    You have been given text extracted from a legal document along with a prompt given by a user. Answer the prompt based on the information given in the text.
    Text:
    {session.get("original")}

    Prompt:
    {user_prompt}
"""
    )

    return res.text


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", use_reloader=True)
